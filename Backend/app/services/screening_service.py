import os
import fitz
import json
import traceback
from docx import Document
from sqlalchemy import func
from sqlalchemy.orm import Session
from langsmith import traceable
from typing import Any
from app.models.resume import Resume
from app.models.screening_result import ScreeningResult
from app.models.job_description import JobDescription

from Agents.Graph import graphs as graph
from app.database.database import SessionLocal


def extract_resume_text(file_path: str):

    extension = os.path.splitext(file_path)[1].lower()

    # ==========================
    # PDF
    # ==========================
    if extension == ".pdf":

        document = fitz.open(file_path)

        text = ""

        for page in document:
            text += page.get_text()

        document.close()

        return text

    # ==========================
    # DOCX / DOC
    # ==========================
    elif extension in [".docx", ".doc"]:

        document = Document(file_path)

        text_chunks = []

        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_chunks.append(paragraph.text.strip())

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    # Remove duplicate adjacent cell values (common in merged docx cells)
                    unique_cells = []
                    for c in row_cells:
                        if not unique_cells or unique_cells[-1] != c:
                            unique_cells.append(c)
                    text_chunks.append(" | ".join(unique_cells))

        return "\n".join(text_chunks)

    # ==========================
    # Unsupported File
    # ==========================
    else:

        raise Exception(
            f"Unsupported file type: {extension}"
        )


# ==============================
# Call LangGraph
# ==============================

@traceable(name="Resume Screening Workflow")
def screen_resume (job_description: str, resume_text: str):

    result = graph.invoke(
        {
            "job_description": job_description,
            "resume": resume_text
        }
    )

    return result   


# ==============================
# Save AI Result
# ==============================
import json

def save_screening_result(
    resume: Resume,
    result: dict,
    db: Session
):

    # =========================================================
    # =========================================================
    # FINAL SCORER OUTPUT
    # =========================================================

    final_score = result.get("final_score", {})

    if isinstance(final_score, str):
        clean_text = final_score.strip()
        if clean_text.startswith("```"):
            lines = clean_text.splitlines()
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            clean_text = "\n".join(lines).strip()
        try:
            final_score = json.loads(clean_text)
        except Exception:
            final_score = {}

    if not isinstance(final_score, dict):
        final_score = {}

    scores = final_score.get("scoring_breakdown", {})
    categories = scores.get("category_scores", {})


    # =========================================================
    # MATCH EVALUATOR OUTPUT
    # =========================================================

    match_evaluation = result.get(
        "match_evaluation",
        {}
    )

    if isinstance(match_evaluation, str):

        try:
            match_evaluation = json.loads(
                match_evaluation
            )

        except Exception:

            match_evaluation = {}


    # =========================================================
    # EXTRACT DETAILED CATEGORY ANALYSIS
    # =========================================================

    category_analysis = {

        "technical_skills": match_evaluation.get(
            "skills_analysis",
            {}
        ),

        "experience": match_evaluation.get(
            "experience_analysis",
            {}
        ),

        "education": match_evaluation.get(
            "education_and_certifications_analysis",
            {}
        ),

        "domain_fit": match_evaluation.get(
            "domain_fit_analysis",
            {}
        )
    }


    # =========================================================
    # DETERMINISTIC RECOMMENDATION FALLBACK
    # =========================================================

    recommendation = (
        final_score.get("hiring_recommendation")
        or final_score.get("recommendation")
        or match_evaluation.get("overall_evaluation", {}).get("hiring_recommendation")
        or match_evaluation.get("overall_evaluation", {}).get("recommendation")
    )

    if not recommendation:
        overall_score = float(scores.get("overall_match_score", 0))
        if overall_score >= 85:
            recommendation = "Strong Hire"
        elif overall_score >= 70:
            recommendation = "Hire"
        elif overall_score >= 55:
            recommendation = "Shortlist"
        elif overall_score >= 40:
            recommendation = "Consider"
        else:
            recommendation = "Reject"


    # =========================================================
    # SAVE SCREENING RESULT
    # =========================================================

    screening = ScreeningResult(

        # -----------------------------------------------------
        # Deterministic scores
        # -----------------------------------------------------

        overall_score=scores.get(
            "overall_match_score", 0
        ),

        technical_score=categories.get(
            "technical_skills_score", 0
        ),

        experience_score=categories.get(
            "experience_score", 0
        ),

        education_score=categories.get(
            "education_score", 0
        ),

        domain_score=categories.get(
            "domain_fit_score", 0
        ),


        # -----------------------------------------------------
        # Deterministic recommendation
        # -----------------------------------------------------

        recommendation=recommendation,


        # -----------------------------------------------------
        # Final Scorer output
        # -----------------------------------------------------

        executive_summary=final_score.get(
            "executive_summary", ""
        ),

        strengths=final_score.get(
            "top_reasons_to_hire", []
        ),

        weaknesses=final_score.get(
            "key_risk_factors", []
        ),

        interview_questions=final_score.get(
            "recommended_interview_questions", []
        ),


        # -----------------------------------------------------
        # Match Evaluator detailed analysis
        # -----------------------------------------------------

        category_analysis=category_analysis,


        # -----------------------------------------------------
        # Resume
        # -----------------------------------------------------

        resume_id=resume.id
    )


    db.add(screening)

    db.commit()

    db.refresh(screening)

    return screening
# ==============================
# Complete Resume Screening Flow
# ==============================
def screen_uploaded_resume(
    resume_id: int,
    db: Session = None
):
    should_close_db = False
    if db is None:
        db = SessionLocal()
        should_close_db = True

    try:
        resume = (
            db.query(Resume)
            .filter(Resume.id == resume_id)
            .first()
        )

        if not resume:
            print("❌ Resume not found")
            return None

        # --------------------
        # Step 1
        # --------------------
        resume.status = "Parsing"
        db.commit()

        print("✅ Status -> Parsing")

        # --------------------
        # Step 2
        # --------------------
        resume_text = extract_resume_text(
            resume.file_path
        )

        print("✅ Resume Text Extracted")
        print(f"Characters : {len(resume_text)}")

        resume.resume_text = resume_text
        db.commit()

        # --------------------
        # Step 3
        # --------------------
        resume.status = "Screening"
        db.commit()

        print("✅ Status -> Screening")

        # --------------------
        # Step 4
        # --------------------
        job = resume.job_description

        print("\nJOB DESCRIPTION")
        print("-" * 80)
        print(job.description[:500])
        print("-" * 80)

        # --------------------
        # Step 5
        # --------------------
        print("\nCalling LangGraph Agent...\n")

        result = screen_resume(
            job.description,
            resume.resume_text
        )

        # --------------------
        # Extract and save candidate info
        # --------------------
        parsed_resume = result.get("parsed_resume")
        if parsed_resume:
            if isinstance(parsed_resume, str):
                try:
                    parsed_resume = json.loads(parsed_resume)
                except Exception:
                    parsed_resume = {}
            if isinstance(parsed_resume, dict):
                candidate = parsed_resume.get("candidate_info", {})
                if candidate:
                    resume.candidate_name = candidate.get("full_name") or candidate.get("name")
                    resume.email = candidate.get("email")
                    resume.phone = candidate.get("phone")
                    db.commit()

        print("\n" + "=" * 80)
        print("AI RESULT")
        print("=" * 80)
        print(result)
        print("TYPE :", type(result))
        print("=" * 80)

        # --------------------
        # Step 6
        # --------------------
        save_screening_result(
            resume,
            result,
            db
        )

        print("✅ AI Result Saved")

        # --------------------
        # Step 7
        # --------------------
        resume.status = "Completed"
        db.commit()

        print("✅ Status -> Completed")

        return result

    except Exception as e:

        resume.status = "Failed"
        db.commit()

        print("\n" + "=" * 80)
        print("❌ SCREENING ERROR")
        print("=" * 80)

        traceback.print_exc()

        print("\nException Message :")
        print(e)

        print("=" * 80)

        return None

    finally:
        if should_close_db:
            db.close()


# ==============================
# Candidate Details
# ==============================
def get_screening_result(
    resume_id: int,
    db: Session
):

    resume = (
        db.query(Resume)
        .filter(
            Resume.id == resume_id
        )
        .first()
    )

    if not resume or not resume.screening_result:
        return None

    res = resume.screening_result

    return {
        "overall_score": res.overall_score,
        "technical_score": res.technical_score,
        "experience_score": res.experience_score,
        "education_score": res.education_score,
        "domain_score": res.domain_score,
        "category_analysis": res.category_analysis or {},
        "recommendation": res.recommendation,
        "executive_summary": res.executive_summary,
        "strengths": res.strengths or [],
        "weaknesses": res.weaknesses or [],
        "interview_questions": res.interview_questions or [],
        "candidate_name": resume.candidate_name or resume.filename,
        "email": resume.email or "N/A",
        "phone": resume.phone or "N/A",
        "filename": resume.filename
    }


# ==============================
# Ranking
# ==============================
def get_job_ranking(
    job_id: int,
    db: Session
):

    resumes = (
        db.query(Resume)
        .filter(
            Resume.job_description_id == job_id,
            Resume.status != "Failed"
        )
        .join(ScreeningResult)
        .order_by(
            ScreeningResult.overall_score.desc()
        )
        .all()
    )

    ranking = []

    for resume in resumes:

        ranking.append({
            "resume_id": resume.id,
            "candidate_name": resume.candidate_name or resume.filename or f"Candidate #{resume.id}",
            "job_title": resume.job_description.title if resume.job_description else "N/A",
            "filename": resume.filename,
            "overall_score": resume.screening_result.overall_score,
            "recommendation": resume.screening_result.recommendation
        })

    return ranking


# ==============================
# Analytics
# ==============================
def get_job_analytics(
    job_id: int,
    db: Session
):

    total = (
        db.query(Resume)
        .filter(
            Resume.job_description_id == job_id
        )
        .count()
    )

    average = (
        db.query(
            func.avg(
                ScreeningResult.overall_score
            )
        )
        .join(Resume)
        .filter(
            Resume.job_description_id == job_id
        )
        .scalar()
    )

    highest = (
        db.query(
            func.max(
                ScreeningResult.overall_score
            )
        )
        .join(Resume)
        .filter(
            Resume.job_description_id == job_id
        )
        .scalar()
    )

    lowest = (
        db.query(
            func.min(
                ScreeningResult.overall_score
            )
        )
        .join(Resume)
        .filter(
            Resume.job_description_id == job_id
        )
        .scalar()
    )

    return {
        "total_candidates": total,
        "average_score": average,
        "highest_score": highest,
        "lowest_score": lowest
    }



def get_all_candidates(db: Session, user_id: Any = None):

    query = (
        db.query(Resume)
        .filter(Resume.status != "Failed")
        .join(ScreeningResult, Resume.id == ScreeningResult.resume_id)
        .join(JobDescription, Resume.job_description_id == JobDescription.id)
    )

    if user_id is not None:
        try:
            uid = int(user_id)
            query = query.filter(JobDescription.created_by == uid)
        except (ValueError, TypeError):
            pass

    resumes = (
        query.order_by(ScreeningResult.overall_score.desc())
        .all()
    )

    candidates = []

    for index, resume in enumerate(resumes, start=1):

        candidates.append({

                "rank": index,

                "resume_id": resume.id,

                "candidate_name": resume.candidate_name or resume.filename,

                "job_title": resume.job_description.title,

                "filename": resume.filename,

                "overall_score": resume.screening_result.overall_score,

                "recommendation": resume.screening_result.recommendation

})

    return candidates